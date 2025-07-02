import streamlit as st
import pandas as pd
from io import BytesIO
import json
from datetime import datetime
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class ReportGenerator:
    """Handles report generation and export functionality"""
    
    def __init__(self):
        pass
    
    def render_report_interface(self, data, dashboard_config):
        """Render the report generation interface"""
        st.subheader("📄 Report Generation")
        
        # Report configuration
        report_config = self._get_report_config(data, dashboard_config)
        
        # Preview section
        if st.button("👀 Preview Report", type="primary"):
            self._preview_report(data, dashboard_config, report_config)
        
        st.markdown("---")
        
        # Export options
        st.subheader("📤 Export Options")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("📊 Export Data (CSV)"):
                self._export_csv(data)
        
        with col2:
            if st.button("📋 Export Summary (JSON)"):
                self._export_json_summary(data, dashboard_config)
        
        with col3:
            if DOCX_AVAILABLE:
                if st.button("📄 Export Report (DOCX)"):
                    self._export_docx_report(data, dashboard_config, report_config)
            else:
                st.button("📄 Export Report (DOCX)", disabled=True, help="python-docx not available")
    
    def _get_report_config(self, data, dashboard_config):
        """Get report configuration from user input"""
        with st.expander("⚙️ Report Configuration", expanded=True):
            config = {}
            
            col1, col2 = st.columns(2)
            
            with col1:
                config['title'] = st.text_input(
                    "Report Title",
                    value="Data Analysis Report"
                )
                
                config['author'] = st.text_input(
                    "Author",
                    value="Analytics Platform"
                )
            
            with col2:
                config['include_summary'] = st.checkbox(
                    "Include Data Summary",
                    value=True
                )
                
                config['include_charts'] = st.checkbox(
                    "Include Charts",
                    value=True
                )
            
            config['description'] = st.text_area(
                "Report Description",
                value="Automated report generated from uploaded data analysis."
            )
            
            return config
    
    def _preview_report(self, data, dashboard_config, report_config):
        """Preview the report content"""
        st.markdown("---")
        st.subheader("📖 Report Preview")
        
        # Report header
        st.markdown(f"# {report_config['title']}")
        st.markdown(f"**Author:** {report_config['author']}")
        st.markdown(f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        st.markdown(f"**Description:** {report_config['description']}")
        
        st.markdown("---")
        
        # Data summary section
        if report_config['include_summary']:
            st.markdown("## 📊 Data Summary")
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.metric("Total Rows", f"{data.shape[0]:,}")
            
            with col2:
                st.metric("Total Columns", data.shape[1])
            
            with col3:
                missing_count = data.isnull().sum().sum()
                st.metric("Missing Values", f"{missing_count:,}")
            
            # Column information
            st.markdown("### Column Information")
            col_info = []
            for col in data.columns:
                col_info.append({
                    'Column': col,
                    'Type': str(data[col].dtype),
                    'Non-Null Count': data[col].count(),
                    'Unique Values': data[col].nunique()
                })
            
            st.dataframe(pd.DataFrame(col_info), use_container_width=True)
            
            # Basic statistics for numeric columns
            numeric_cols = data.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                st.markdown("### Numeric Columns Statistics")
                st.dataframe(data[numeric_cols].describe(), use_container_width=True)
        
        # Dashboard components section
        if report_config['include_charts'] and dashboard_config:
            st.markdown("## 📈 Visualizations")
            
            from utils.visualization import Visualization
            viz = Visualization()
            
            for i, component in enumerate(dashboard_config):
                st.markdown(f"### {component.get('title', f'Component {i+1}')}")
                
                if component['type'] == 'chart':
                    try:
                        fig = viz.create_chart(
                            component['chart_type'],
                            data,
                            component['config']
                        )
                        if fig:
                            st.plotly_chart(fig, use_container_width=True)
                    except Exception as e:
                        st.error(f"Error displaying chart: {str(e)}")
                
                elif component['type'] == 'metric':
                    config = component['config']
                    try:
                        column = config.get('column')
                        aggregation = config.get('aggregation', 'sum')
                        
                        if column and column in data.columns:
                            if aggregation == 'sum':
                                value = data[column].sum()
                            elif aggregation == 'mean':
                                value = data[column].mean()
                            elif aggregation == 'count':
                                value = data[column].count()
                            elif aggregation == 'min':
                                value = data[column].min()
                            elif aggregation == 'max':
                                value = data[column].max()
                            elif aggregation == 'median':
                                value = data[column].median()
                            
                            st.metric(component['title'], f"{value:,.2f}")
                    except Exception as e:
                        st.error(f"Error displaying metric: {str(e)}")
                
                elif component['type'] == 'table':
                    config = component['config']
                    try:
                        selected_cols = config.get('columns', [])
                        max_rows = config.get('max_rows', 100)
                        
                        if selected_cols:
                            display_data = data[selected_cols].head(max_rows)
                            st.dataframe(display_data, use_container_width=True)
                    except Exception as e:
                        st.error(f"Error displaying table: {str(e)}")
    
    def _export_csv(self, data):
        """Export data as CSV"""
        try:
            csv_buffer = BytesIO()
            data.to_csv(csv_buffer, index=False)
            csv_buffer.seek(0)
            
            st.download_button(
                label="⬇️ Download CSV",
                data=csv_buffer.getvalue(),
                file_name=f"data_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv"
            )
            
            st.success("✅ CSV export ready for download!")
            
        except Exception as e:
            st.error(f"❌ Error exporting CSV: {str(e)}")
    
    def _export_json_summary(self, data, dashboard_config):
        """Export summary information as JSON"""
        try:
            summary = {
                'export_info': {
                    'generated_at': datetime.now().isoformat(),
                    'platform': 'Analytics Platform v1.0'
                },
                'data_summary': {
                    'shape': data.shape,
                    'columns': data.columns.tolist(),
                    'dtypes': data.dtypes.astype(str).to_dict(),
                    'missing_values': data.isnull().sum().to_dict(),
                    'numeric_summary': data.describe().to_dict() if len(data.select_dtypes(include=['number']).columns) > 0 else {}
                },
                'dashboard_config': dashboard_config
            }
            
            json_str = json.dumps(summary, indent=2, default=str)
            
            st.download_button(
                label="⬇️ Download JSON Summary",
                data=json_str,
                file_name=f"analysis_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json"
            )
            
            st.success("✅ JSON summary export ready for download!")
            
        except Exception as e:
            st.error(f"❌ Error exporting JSON: {str(e)}")
    
    def _export_docx_report(self, data, dashboard_config, report_config):
        """Export comprehensive report as DOCX"""
        if not DOCX_AVAILABLE:
            st.error("❌ python-docx library not available for DOCX export")
            return
        
        try:
            # Create document
            doc = Document()
            
            # Add title
            title = doc.add_heading(report_config['title'], 0)
            
            # Add metadata
            doc.add_paragraph(f"Author: {report_config['author']}")
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            doc.add_paragraph(f"Description: {report_config['description']}")
            
            # Add data summary
            if report_config['include_summary']:
                doc.add_heading('Data Summary', level=1)
                
                # Basic statistics
                doc.add_paragraph(f"Dataset contains {data.shape[0]:,} rows and {data.shape[1]} columns.")
                
                missing_count = data.isnull().sum().sum()
                doc.add_paragraph(f"Total missing values: {missing_count:,}")
                
                # Column information
                doc.add_heading('Column Information', level=2)
                
                # Create table for column info
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                
                # Add header row
                header_cells = table.rows[0].cells
                header_cells[0].text = 'Column'
                header_cells[1].text = 'Type'
                header_cells[2].text = 'Non-Null Count'
                header_cells[3].text = 'Unique Values'
                
                # Add data rows
                for col in data.columns:
                    row_cells = table.add_row().cells
                    row_cells[0].text = col
                    row_cells[1].text = str(data[col].dtype)
                    row_cells[2].text = str(data[col].count())
                    row_cells[3].text = str(data[col].nunique())
            
            # Add dashboard information
            if dashboard_config:
                doc.add_heading('Dashboard Components', level=1)
                
                for i, component in enumerate(dashboard_config):
                    doc.add_heading(f"{component.get('title', f'Component {i+1}')}", level=2)
                    doc.add_paragraph(f"Type: {component['type']}")
                    
                    if component['type'] == 'chart':
                        doc.add_paragraph(f"Chart Type: {component.get('chart_type', 'N/A')}")
                        doc.add_paragraph(f"Configuration: {json.dumps(component.get('config', {}), indent=2)}")
            
            # Save to BytesIO
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            st.download_button(
                label="⬇️ Download DOCX Report",
                data=doc_buffer.getvalue(),
                file_name=f"analysis_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            st.success("✅ DOCX report export ready for download!")
            
        except Exception as e:
            st.error(f"❌ Error exporting DOCX: {str(e)}")
